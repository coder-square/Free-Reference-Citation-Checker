"""
Tests for reference_parser.py. No network needed -- these run entirely
against text fixtures and real .docx files built with python-docx, so
every test here is a genuine end-to-end check, not a mock.
"""

import os
import tempfile

import pytest
from docx import Document

from reference_parser import (
    split_references,
    extract_text_from_docx,
)


# ---------- split_references: blank-line strategy ----------

BLANK_LINE_LIST = """Fasuan, T. M., & Adebayo, O. (2023). Okra mucilage films for postharvest fruit coating. Journal of Postharvest Science, 8(2), 45-58.

Bello, K., & Smith, P. (2022). Non-thermal pasteurization of passion fruit juice. Food Engineering Reviews, 14(1), 12-30.

Olagunju, A. (2021). Starch composition in composite flours. International Journal of Food Science, 9(3), 200-215."""

def test_split_by_blank_lines():
    result = split_references(BLANK_LINE_LIST)
    assert result.method == "blank_line"
    assert len(result.entries) == 3
    assert "Fasuan" in result.entries[0]
    assert "Bello" in result.entries[1]
    assert "Olagunju" in result.entries[2]

def test_split_by_blank_lines_ignores_stray_blank_paragraphs():
    # A single accidental blank line inside otherwise unstructured text
    # should not be mistaken for a real reference-list split.
    text = "This is just a short note.\n\nNot a reference at all."
    result = split_references(text)
    assert result.method != "blank_line"


# ---------- split_references: numbered strategy ----------

NUMBERED_LIST = (
    "1. Fasuan, T. M. (2023). Okra mucilage films. Journal of Postharvest Science, 8(2), 45-58.\n"
    "2. Bello, K. (2022). Non-thermal pasteurization. Food Engineering Reviews, 14(1), 12-30.\n"
    "3. Olagunju, A. (2021). Starch composition. Int J Food Sci, 9(3), 200-215.\n"
)

def test_split_by_numbering():
    result = split_references(NUMBERED_LIST)
    assert result.method == "numbered"
    assert len(result.entries) == 3
    assert result.entries[0].startswith("Fasuan")

def test_split_by_bracketed_numbering():
    text = "[1] Fasuan T (2023). Title one.\n[2] Bello K (2022). Title two.\n"
    result = split_references(text)
    assert result.method == "numbered"
    assert len(result.entries) == 2


# ---------- split_references: author-pattern fallback ----------

def test_split_by_author_pattern_fallback():
    # No blank lines, no numbers -- continuous text with author-year boundaries.
    text = (
        "Fasuan, T. Okra mucilage films for coating. J Postharvest Sci 2023. "
        "Bello, K. Non-thermal pasteurization methods. Food Eng Rev 2022. "
        "Olagunju, A. Starch composition analysis. Int J Food Sci 2021."
    )
    result = split_references(text)
    assert result.method == "author_pattern"
    assert len(result.entries) == 3
    assert result.warning  # must warn that this is the less reliable method


# ---------- split_references: honest failure case ----------

def test_split_unparseable_text_returns_single_block_with_warning():
    text = "just some random text with no structure at all and no year 20xx pattern"
    result = split_references(text)
    assert result.method == "single_block"
    assert result.entries == [text]
    assert result.warning  # must not silently pretend this worked

def test_split_empty_text():
    result = split_references("")
    assert result.method == "empty"
    assert result.entries == []

def test_split_whitespace_only():
    result = split_references("   \n\n   ")
    assert result.method == "empty"


# ---------- extract_text_from_docx: real files, no mocks ----------

def _make_docx(paragraphs: list) -> str:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    return path

def test_extract_docx_finds_references_heading_and_skips_body():
    path = _make_docx([
        "Introduction",
        "This manuscript discusses okra mucilage coatings in detail.",
        "Methods",
        "We used a randomized design.",
        "References",
        "Fasuan, T. M. (2023). Okra mucilage films. Journal of Postharvest Science, 8(2), 45-58.",
        "Bello, K. (2022). Non-thermal pasteurization. Food Engineering Reviews, 14(1), 12-30.",
    ])
    try:
        result = extract_text_from_docx(path)
        assert result.found_heading is True
        assert result.paragraph_count == 2
        assert "Introduction" not in result.text
        assert "randomized design" not in result.text
        assert "Fasuan" in result.text
        assert "Bello" in result.text
        assert result.warning == ""
    finally:
        os.remove(path)

def test_extract_docx_no_heading_returns_everything_with_warning():
    path = _make_docx([
        "Fasuan, T. M. (2023). Okra mucilage films. Journal of Postharvest Science, 8(2), 45-58.",
        "Bello, K. (2022). Non-thermal pasteurization. Food Engineering Reviews, 14(1), 12-30.",
    ])
    try:
        result = extract_text_from_docx(path)
        assert result.found_heading is False
        assert result.paragraph_count == 2
        assert result.warning  # must flag that it couldn't confirm scope
    finally:
        os.remove(path)

def test_extract_docx_recognizes_bibliography_heading_variant():
    path = _make_docx([
        "Some intro text.",
        "Bibliography",
        "Fasuan, T. M. (2023). Okra mucilage films. Journal of Postharvest Science, 8(2), 45-58.",
    ])
    try:
        result = extract_text_from_docx(path)
        assert result.found_heading is True
        assert result.paragraph_count == 1
    finally:
        os.remove(path)

def test_extract_docx_ignores_blank_paragraphs():
    doc = Document()
    doc.add_paragraph("References")
    doc.add_paragraph("")  # blank paragraph, common in real Word docs
    doc.add_paragraph("Fasuan, T. M. (2023). Okra mucilage films. Journal of Postharvest Science, 8(2), 45-58.")
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    try:
        result = extract_text_from_docx(path)
        assert result.paragraph_count == 1  # blank paragraph must not count as an entry
    finally:
        os.remove(path)


# ---------- end-to-end: docx extraction feeding straight into the splitter ----------

def test_docx_extraction_feeds_splitter_correctly():
    path = _make_docx([
        "References",
        "Fasuan, T. M. (2023). Okra mucilage films for postharvest fruit coating. Journal of Postharvest Science, 8(2), 45-58.",
        "Bello, K., & Smith, P. (2022). Non-thermal pasteurization of passion fruit juice. Food Engineering Reviews, 14(1), 12-30.",
        "Olagunju, A. (2021). Starch composition in composite flours. International Journal of Food Science, 9(3), 200-215.",
    ])
    try:
        extraction = extract_text_from_docx(path)
        split = split_references(extraction.text)
        # One reference per paragraph in the docx should mean the blank-line
        # strategy fires -- the most reliable path -- since paragraphs are
        # joined with blank lines by extract_text_from_docx.
        assert split.method == "blank_line"
        assert len(split.entries) == 3
    finally:
        os.remove(path)
